# complete_rag_test.py
"""
完整的RAG测试程序 - 使用您的Qwen2.5配置

RAG (Retrieval-Augmented Generation) 检索增强生成
工作流程:
1. 文档加载 -> 2. 文本分块 -> 3. 向量化存储 -> 4. 相似度检索 -> 5. 构建提示词 -> 6. LLM生成回答
"""

# 导入标准库
import os
import shutil
from pathlib import Path
from typing import List

# 导入第三方库
from docx import Document  # 用于创建Word文档
from langchain_openai import ChatOpenAI  # OpenAI兼容的LLM接口（用于连接Ollama）
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader  # 文档加载器
from langchain_text_splitters import RecursiveCharacterTextSplitter  # 文本分块器
from langchain_community.vectorstores import Chroma  # 向量数据库
from langchain_huggingface import HuggingFaceEmbeddings  # 嵌入模型（将文本转为向量）
from langchain_core.messages import HumanMessage, SystemMessage  # 消息类型定义
from langchain_openai import OpenAIEmbeddings
from langchain_core.embeddings import Embeddings
import requests
from typing import List


class LocalRAG:
    """
    本地RAG系统核心类
    
    功能：
    - 文档加载与向量化存储
    - 相似度检索
    - 基于检索结果的回答生成
    
    核心组件：
    - LLM: 大语言模型（连接本地Ollama的Qwen2.5）
    - embeddings: 嵌入模型（将文本转为向量）
    - vector_store: 向量数据库（Chroma）
    - text_splitter: 文本分块器
    """
    
    def __init__(self, 
                 llm_base_url: str = "http://192.168.3.8:11434/v1",
                 llm_api_key: str = "ollama",
                 model_name: str = "qwen2.5:14b-custom",
                 embedding_model: str = "nomic-embed-text",
                 persist_directory: str = "./chroma_db"):
        
        # ========== 1. 初始化大语言模型 (LLM) ==========
        # 使用ChatOpenAI接口连接Ollama服务
        # 注意：Ollama通过v1兼容API提供OpenAI格式接口
        self.llm = ChatOpenAI(
            model=model_name,  # 模型名称
            base_url=llm_base_url,  # Ollama服务地址
            api_key=llm_api_key,  # API密钥（Ollama用ollama）
            temperature=0.1,  # 温度参数：0.1表示输出较确定，减少随机性
            timeout=120,  # 超时时间：120秒
        )
        self.model_name = model_name
        
        print("正在加载嵌入模型...")
        
        # ========== 2. 初始化嵌入模型 (Embeddings) ==========
        # 嵌入模型的作用：将文本转换为向量
        # 这样可以在向量空间中计算文本之间的相似度
        # BAAI/bge-small-zh-v1.5: 轻量级中文嵌入模型
        # self.embeddings = HuggingFaceEmbeddings(
        #     model_name=embedding_model,
        #     model_kwargs={'device': 'cuda'},  # 使用GPU加速（有GPU时设为'cuda'，无GPU时用'cpu'）
        #     encode_kwargs={'normalize_embeddings': True}  # 归一化向量，使相似度计算更稳定
        # )
        self.embeddings = OllamaEmbeddings(
            base_url="http://192.168.3.8:11434",  # 注意：不要 /v1
            model="nomic-embed-text"
        )
        
        # ========== 3. 初始化向量数据库 (Chroma) ==========
        # Chroma是一个轻量级的向量数据库，用于存储文档向量
        # persist_directory: 持久化存储路径，重启后数据不丢失
        self.vector_store = Chroma(
            persist_directory=persist_directory,
            embedding_function=self.embeddings  # 绑定嵌入模型
        )
        
        # ========== 4. 初始化文本分块器 ==========
        # 将长文档切分成小块，便于检索和嵌入
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  # 每个块最大500字符
            chunk_overlap=50,  # 块之间重叠50字符，保持上下文连贯
            # 按优先级分割：段落 -> 句子 -> 标点
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]
        )
        
        print("RAG系统初始化完成！")
    
    def load_documents(self, folder_path: str):
        """
        加载文件夹中的所有Word和PDF文档
        
        处理流程：
        1. 扫描文件夹查找PDF和Word文件
        2. 使用对应加载器读取文档内容
        3. 将长文档切分成小块（chunk）
        4. 将文本块转换为向量并存储到Chroma数据库
        
        Args:
            folder_path: 文档所在文件夹路径
            
        Returns:
            加载的文本块数量
        """
        documents = []  # 存储加载的文档
        folder = Path(folder_path)
        
        # 扫描查找PDF和Word文件
        pdf_files = list(folder.glob("*.pdf"))
        docx_files = list(folder.glob("*.docx"))
        
        print(f"找到 {len(pdf_files)} 个PDF文件, {len(docx_files)} 个Word文件")
        
        # ========== 加载PDF文件 ==========
        for pdf_path in pdf_files:
            try:
                # PyPDFLoader: 专门用于加载PDF文档
                loader = PyPDFLoader(str(pdf_path))
                docs = loader.load()  # load()返回Document列表
                # 为每个文档添加来源元数据，便于后续追溯
                for doc in docs:
                    doc.metadata["source"] = pdf_path.name
                documents.extend(docs)  # 添加到文档列表
                print(f"✓ 已加载: {pdf_path.name}")
            except Exception as e:
                print(f"✗ 加载失败 {pdf_path.name}: {e}")
        
        # ========== 加载Word文件 ==========
        for docx_path in docx_files:
            try:
                # Docx2txtLoader: 专门用于加载Word文档
                loader = Docx2txtLoader(str(docx_path))
                docs = loader.load()
                for doc in docs:
                    doc.metadata["source"] = docx_path.name
                documents.extend(docs)
                print(f"✓ 已加载: {docx_path.name}")
            except Exception as e:
                print(f"✗ 加载失败 {docx_path.name}: {e}")
        
        # 如果没有找到任何文档
        if not documents:
            print("未找到任何文档！")
            return 0
        
        # ========== 文本分块 ==========
        # 将长文档切分成小块，便于后续向量检索
        # chunk_size=500: 每块约500字符
        # chunk_overlap=50: 块之间有50字符重叠，防止信息丢失
        chunks = self.text_splitter.split_documents(documents)
        print(f"已生成 {len(chunks)} 个文本块")
        
        # ========== 向量化存储 ==========
        # 将文本块转换为向量并存储到Chroma数据库
        print("正在构建向量索引...")
        self.vector_store.add_documents(chunks)  # 添加到数据库
        self.vector_store.persist()  # 持久化保存到磁盘
        print("向量数据库构建完成！")
        
        return len(chunks)
    
    def search(self, query: str, top_k: int = 3):
        """
        检索相关文档片段（核心检索方法）
        
        工作原理：
        1. 将用户问题转换为向量（使用embedding模型）
        2. 在向量数据库中查找最相似的文档块
        3. 返回top_k个最相关的文档片段
        
        Args:
            query: 用户问题
            top_k: 返回最相似的k个文档块（默认3个）
            
        Returns:
            包含文档内容、来源和相似度分数的列表
        """
        # similarity_search_with_score: 带分数的相似度搜索
        # 返回：(Document对象, 相似度分数)
        # 分数越低表示越相似（Chroma使用余弦距离）
        results = self.vector_store.similarity_search_with_score(query, k=top_k)
        
        contexts = []
        for doc, score in results:
            contexts.append({
                "content": doc.page_content,  # 文档内容
                "source": doc.metadata.get("source", "unknown"),  # 来源文件
                "score": score  # 相似度分数
            })
        return contexts
    
    def ask(self, question: str, top_k: int = 3, show_context: bool = False) -> str:
        """
        基于RAG回答问题（完整流程）
        
        RAG回答流程：
        1. 检索(Retrieval): 从向量数据库中找到相关文档
        2. 增强(Augmentation): 将检索到的文档拼接到提示词
        3. 生成(Generation): 调用LLM基于文档生成回答
        
        Args:
            question: 用户问题
            top_k: 检索的文档块数量
            show_context: 是否打印检索到的文档（用于调试）
            
        Returns:
            LLM生成的回答文本
        """
        # ========== 步骤1: 检索相关文档 ==========
        contexts = self.search(question, top_k)
        
        if not contexts:
            return "未找到相关文档，请先上传文档。"
        
        # ========== 步骤2: 构建提示词 (Prompt Engineering) ==========
        # 将检索到的文档格式化成上下文
        context_text = "\n\n---\n\n".join([
            f"【文档来源：{ctx['source']}】\n{ctx['content']}" 
            for ctx in contexts
        ])
        
        # 构建消息列表：系统消息 + 用户消息
        # SystemMessage: 定义LLM的角色和行为规则
        # HumanMessage: 用户的问题和提供的内容
        messages = [
            SystemMessage(content="你是一个基于文档回答问题的助手。只使用提供的信息，不编造内容。如果文档中没有相关信息，直接说'根据现有文档无法回答这个问题'。"),
            HumanMessage(content=f"""请基于下面的【参考文档】回答问题：

【参考文档】
{context_text}

【问题】
{question}

要求：
1. 只基于参考文档回答
2. 如果文档中没有答案，直接说无法回答
3. 回答时要引用文档来源
4. 不要编造文档中没有的内容

【回答】""")
        ]
        
        # ========== 步骤3: 调用Qwen模型生成回答 ==========
        try:
            # invoke: 调用LLM处理消息并获取回复
            response = self.llm.invoke(messages)
            answer = response.content
            
            # 可选：显示检索到的上下文（调试用）
            if show_context:
                print("\n" + "="*50)
                print("检索到的相关文档片段：")
                for i, ctx in enumerate(contexts, 1):
                    # 相关度 = 1 - 分数（因为分数是距离，越小越相似）
                    print(f"\n[{i}] 来源：{ctx['source']} (相关度: {1-ctx['score']:.2f})")
                    print(f"内容：{ctx['content'][:200]}...")
                print("="*50 + "\n")
            
            return answer
            
        except Exception as e:
            return f"调用模型失败: {e}"
    
    def clear_database(self):
        """
        清空向量数据库
        
        用于：
        - 重新开始构建知识库
        - 删除错误的向量数据
        - 释放磁盘空间
        """
        # 检查数据库目录是否存在
        if os.path.exists("./chroma_db"):
            # shutil.rmtree: 递归删除目录及其内容
            shutil.rmtree("./chroma_db")
            print("已清空数据库")
            # 重新初始化空的Chroma数据库
            self.vector_store = Chroma(
                persist_directory="./chroma_db",
                embedding_function=self.embeddings
            )

class OllamaEmbeddings(Embeddings):
    """适配 Ollama 的 Embeddings 类"""
    
    def __init__(self, base_url: str = "http://192.168.3.8:11434", model: str = "nomic-embed-text"):
        self.base_url = base_url.rstrip('/')
        self.model = model
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """批量嵌入文档"""
        embeddings = []
        for text in texts:
            embeddings.append(self._embed(text))
        return embeddings
    
    def embed_query(self, text: str) -> List[float]:
        """嵌入单个查询"""
        return self._embed(text)
    
    def _embed(self, text: str) -> List[float]:
        """调用 Ollama API 获取向量"""
        url = f"{self.base_url}/api/embeddings"
        payload = {
            "model": self.model,
            "prompt": text
        }
        
        try:
            response = requests.post(url, json=payload)
            response.raise_for_status()
            data = response.json()
            return data['embedding']
        except Exception as e:
            print(f"Embedding 调用失败: {e}")
            raise

def create_test_document():
    """
    创建一个测试文档（内容模型肯定不知道）
    
    这个文档用来测试RAG系统的核心价值：
    - 模型训练数据中没有这些信息
    - 只有通过RAG检索文档才能正确回答
    - 验证"检索增强"是否真正起作用
    
    文档内容包含：
    - 项目名称：极光项目
    - 负责人、预算、团队成员等内部信息
    - 一些具体的数据（响应速度、QPS、成本节约等）
    """
    # 确保目录存在
    os.makedirs("./documents", exist_ok=True)
    
    # 创建Word文档对象
    doc = Document()
    
    # 添加标题
    doc.add_heading('极光项目内部报告', 0)
    
    # 添加文档内容（段落）
    doc.add_paragraph('项目启动日期：2026年1月15日')
    doc.add_paragraph('项目负责人：张明华')
    doc.add_paragraph('预算金额：2,380,000元')
    doc.add_paragraph('实际完成日期：2026年3月28日')
    doc.add_paragraph('团队成员：王丽、李思思、陈伟、赵芳')
    doc.add_paragraph('响应速度优化：从200ms降至45ms')
    doc.add_paragraph('并发处理峰值：5800 QPS')
    doc.add_paragraph('成本节约：42.7%')
    doc.add_paragraph('遗留问题：数据库迁移时出现3次连接超时')
    doc.add_paragraph('下季度计划：5月1日上线移动端适配版本')
    
    # 保存文档
    doc.save('./documents/极光项目报告.docx')
    print("✓ 已创建测试文档：极光项目报告.docx")


def test_pure_model(llm: ChatOpenAI):
    """
    测试纯模型回答（不提供文档）
    
    对比实验：
    - 不使用RAG，直接问模型问题
    - 验证模型是否知道"极光项目"的信息
    - 证明模型确实不知道这些内部信息
    
    预期结果：
    - 模型无法回答或回答错误（因为训练数据中没有这些内容）
    """
    print("\n" + "="*60)
    print("【测试1】纯Qwen模型回答（无文档支持）")
    print("="*60)
    
    # 测试问题列表 - 都是文档中有的信息
    questions = [
        "极光项目的负责人是谁？",
        "极光项目的预算是多少钱？",
        "响应速度从多少优化到了多少？",
        "成本节约了百分之多少？",
    ]
    
    for q in questions:
        try:
            # 直接调用LLM，不提供任何上下文
            response = llm.invoke([HumanMessage(content=q)])
            answer = response.content
            print(f"\n问题：{q}")
            print(f"纯模型回答：{answer[:150]}...")
        except Exception as e:
            print(f"错误：{e}")


def test_rag_system(rag: LocalRAG):
    """
    测试RAG系统回答（基于文档）
    
    测试内容：
    1. 文档中有的信息：RAG能否正确回答
    2. 文档中没有的信息：RAG能否拒绝回答
    
    测试用例设计：
    - 前4个问题：文档中有答案
    - 后2个问题：文档中没有答案（测试模型的拒绝能力）
    """
    print("\n" + "="*60)
    print("【测试2】RAG系统回答（基于文档）")
    print("="*60)
    
    # 测试用例：问题 + 预期结果
    test_cases = [
        {"question": "极光项目的负责人是谁？", "expected": "张明华"},
        {"question": "极光项目的预算是多少钱？", "expected": "2,380,000元"},
        {"question": "响应速度从多少优化到了多少？", "expected": "200ms到45ms"},
        {"question": "成本节约了百分之多少？", "expected": "42.7%"},
        {"question": "极光项目是哪个公司做的？", "expected": "拒绝回答（文档没有）"},
        {"question": "今天的天气怎么样？", "expected": "拒绝回答（文档没有）"},
    ]
    
    for tc in test_cases:
        print(f"\n问题：{tc['question']}")
        print(f"预期：{tc['expected']}")
        answer = rag.ask(tc['question'], show_context=False)
        print(f"RAG回答：{answer}")
        print("-"*40)


def compare_answers(rag: LocalRAG, llm: ChatOpenAI):
    """
    直接对比两种方式的回答
    
    对比实验设计：
    - 同一问题分别用"纯模型"和"RAG系统"回答
    - 观察差异：RAG能否基于文档给出准确答案
    
    预期结果：
    - 纯模型：不知道"极光项目"的信息（训练数据中没有）
    - RAG系统：能准确回答（因为检索到了文档中的信息）
    """
    print("\n" + "="*60)
    print("【测试3】直接对比：纯模型 vs RAG")
    print("="*60)
    
    question = "极光项目的负责人是谁？"
    
    # 方法1: 纯模型（不使用RAG）
    response = llm.invoke([HumanMessage(content=question)])
    pure_answer = response.content
    
    # 方法2: RAG系统（使用文档检索）
    rag_answer = rag.ask(question)
    
    print(f"\n问题：{question}")
    print(f"\n📝 纯Qwen模型：{pure_answer}")
    print(f"\n📚 RAG系统：{rag_answer}")
    
    # 分析结果
    if "张明华" in rag_answer and "张明华" not in pure_answer:
        print("\n💡 结论：RAG系统基于文档给出了准确答案，纯模型不知道这个信息")
    elif "张明华" in rag_answer and "张明华" in pure_answer:
        print("\n💡 结论：两者都答对了，但RAG引用了文档来源")
    else:
        print("\n💡 结论：请检查RAG是否正确加载了文档")


def interactive_mode(rag: LocalRAG):
    """
    交互式测试模式
    
    功能：
    - 让用户输入问题进行测试
    - 实时查看RAG的检索和回答效果
    - 输入'exit'退出
    
    使用场景：
    - 调试RAG系统
    - 测试特定问题的回答效果
    - 验证文档检索的准确性
    """
    print("\n" + "="*60)
    print("【交互模式】输入问题测试RAG效果（输入 'exit' 退出）")
    print("💡 提示：问一些文档里的内容，看模型是否准确回答")
    print("="*60)
    
    # 持续接收用户输入
    while True:
        question = input("\n您的问题: ").strip()
        # 检查退出命令
        if question.lower() in ['exit', 'quit', 'q']:
            break
        if not question:
            continue
        
        print("\n正在检索文档...")
        answer = rag.ask(question, show_context=True)
        print(f"\n🤖 回答：{answer}")


def main():
    """
    主函数 - RAG系统完整测试流程
    
    执行步骤：
    1. 配置Ollama连接参数
    2. 测试Ollama服务连接
    3. 创建测试文档
    4. 初始化RAG系统
    5. 加载文档到向量数据库
    6. 运行各项测试：
       - 测试1: 纯模型回答
       - 测试2: RAG系统回答
       - 测试3: 对比两种方式
       - 测试4: 交互模式
    """
    # ========== 使用您的配置 ==========
    OLLAMA_BASE_URL = "http://192.168.3.8:11434/v1"
    MODEL_NAME = "qwen2.5:14b-custom"
    # ============================
    
    print("\n🚀 启动RAG完整测试程序")
    print(f"📡 Ollama服务地址：{OLLAMA_BASE_URL}")
    print(f"🤖 模型名称：{MODEL_NAME}")
    
    # ========== 测试连接 ==========
    print("\n🔌 测试Ollama连接...")
    test_llm = ChatOpenAI(
        model=MODEL_NAME,
        base_url=OLLAMA_BASE_URL,
        api_key="ollama",
        temperature=0.1,
        timeout=120,
    )
    
    try:
        # 发送简单的测试消息验证连接
        test_response = test_llm.invoke([HumanMessage(content="你好，请简单回复'连接成功'")])
        print(f"✓ 连接成功！模型回复：{test_response.content[:50]}")
    except Exception as e:
        print(f"✗ 连接失败：{e}")
        print("请检查：")
        print("1. Ollama服务是否在 192.168.3.8:11434 上运行")
        print("2. 模型 qwen2.5:14b-custom 是否已安装")
        print("3. 网络是否可以访问该地址")
        return
    
    # ========== 步骤1: 创建测试文档 ==========
    print("\n📝 步骤1：创建测试文档...")
    create_test_document()
    
    # ========== 步骤2: 初始化RAG系统 ==========
    print("\n🔧 步骤2：初始化RAG系统...")
    rag = LocalRAG(
        llm_base_url=OLLAMA_BASE_URL,
        llm_api_key="ollama",
        model_name=MODEL_NAME
    )
    
    # ========== 步骤3: 加载文档到RAG ==========
    print("\n📚 步骤3：加载文档到RAG...")
    rag.load_documents("./documents")
    
    # ========== 步骤4: 运行测试 ==========
    print("\n" + "🎯"*30)
    print("开始测试验证")
    print("🎯"*30)
    
    # 测试1：纯模型（不使用RAG）
    test_pure_model(test_llm)
    
    # 测试2：RAG系统（使用文档）
    test_rag_system(rag)
    
    # 测试3：直接对比
    compare_answers(rag, test_llm)
    
    # 测试4：交互模式
    interactive_mode(rag)


if __name__ == "__main__":
    print("请确保已安装依赖：")
    print("pip install langchain langchain-community chromadb sentence-transformers pypdf python-docx langchain-openai")
    print("\n如果已安装，按回车继续...")
    input()
    
    main()